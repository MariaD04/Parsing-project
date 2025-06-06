from docx import Document
import httpx
import logging
import re
from typing import Optional
from openpyxl import Workbook, load_workbook


class ParsingTask:
   
    def __init__(self, docx_in: str, xlsx_in: str, docx_out: str, xlsx_out: str,
        ollama_url: str = "http://localhost:11434/api/generate",
        model: str = "gemma:2b",  
        timeout: float = 60.0
    ):
        self.docx_in = docx_in
        self.xlsx_in = xlsx_in
        self.docx_out = docx_out
        self.xlsx_out = xlsx_out
        self.ollama_url = ollama_url
        self.model = model
        self.timeout = timeout
        self.logger = logging.getLogger(self.__class__.__name__)
   
    async def query_ollama(self, prompt: str) -> Optional[dict]:
        """Отправление запроса к Ollama и получение ответа"""
        try:
            async with httpx.AsyncClient(timeout=self.timeout) as client:
                response = await client.post(
                    self.ollama_url,
                    json={
                        "model": self.model,
                        "prompt": prompt,
                        "stream": False
                    }
                )
                response.raise_for_status()
                return response.json()
        except Exception as e:
            self.logger.error(f"Error querying Ollama: {str(e)}")
            raise

    def extract_text_from_docx(self) -> str:
        """Извлечение текста из DOCX файла"""
        try:
            doc = Document(self.docx_in)
            full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
            return full_text
        except Exception as e:
            self.logger.error(f"Error reading DOCX file: {str(e)}")
            raise

    def save_summary_to_docx(self, summary: str) -> None:
        """Сохранение результата в DOCX файл"""
        try:
            out_doc = Document()
            out_doc.add_paragraph(summary)
            out_doc.save(self.docx_out)
        except Exception as e:
            self.logger.error(f"Error saving DOCX file: {str(e)}")
            raise
    
    def analyze_xlsx(self) -> dict:
        """Извлечение текста, количества символов и количества цифр из XLSX файла"""
        try:
            wb = load_workbook(self.xlsx_in)
            all_text = ""
            
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    for cell in row:
                        if cell is not None:
                            all_text += str(cell) + " "
            
            count_char = len(all_text)
            digit_count = len(re.findall(r'\d', all_text))

            return { "all_text": all_text, "count_char": count_char, "digit_count": digit_count }
        
        except Exception as e:
            self.logger.error(f"Error reading or processing XLSX file: {str(e)}")
            raise

    def save_all_to_xlsx(self, all_text_xlsx, count_char_xlsx, count_digits_xlsx):
        """Создание и заполнение нового XLSX файла с результатами"""
        out_wb = Workbook()
        ws = out_wb.active
        
        ws.append([f"Имя файла: {self.xlsx_in}"])
        ws.append([f"Общее количество символов: {count_char_xlsx}"])
        ws.append([f"Общее количество цифр: {count_digits_xlsx}"])
        ws.append([""])  # Пустая строка
        
        if isinstance(all_text_xlsx, str):
            for line in all_text_xlsx.split('\n'):
                ws.append([line])
        
        out_wb.save(self.xlsx_out)

    async def run(self) -> dict:
        """Запуск задачи"""
        try:
            # 1. Обработка DOCX файла
            all_text_docx = self.extract_text_from_docx()
            self.logger.info(f"Extracted text from {self.docx_in}")

            # Генерация краткого содержания
            ollama_response_docx = await self.query_ollama(f"Сделай краткое содержание следующего текста:\n{all_text_docx}")
            summary_docx = ollama_response_docx.get("response", "<Не удалось получить ответ от модели>")
            
            # Сохранение результата в DOCX файл
            self.save_summary_to_docx(summary_docx)
            self.logger.info(f"Summary saved to {self.docx_out}")

            # 2. Обработка XLSX файла
            analyze_results = self.analyze_xlsx()
            all_text_xlsx = analyze_results['all_text']
            count_char_xlsx = analyze_results['count_char']
            count_digits_xlsx = analyze_results['digit_count']
            self.logger.info(f"XLSX analysis saved to {self.xlsx_out}")

             # Генерация краткого содержания
            ollama_response_xlsx = await self.query_ollama(f"Сделай краткое содержание следующего текста:\n{all_text_xlsx}")
            summary_xlsx = ollama_response_xlsx.get("response", "<Не удалось получить ответ от модели>")

            # Сохранение результата в XLSX файл
            self.save_all_to_xlsx(summary_xlsx, count_char_xlsx, count_digits_xlsx)
            self.logger.info(f"Summary saved to {self.docx_out}")

            return {
                "docx_summary": summary_docx,
                "xlsx_summary": summary_xlsx
            }
        except Exception as e:
            self.logger.error(f"Error in parsing task: {str(e)}")
            raise







